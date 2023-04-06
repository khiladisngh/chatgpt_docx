import os
import openai
import docx
import time
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from termcolor import colored
from dotenv import load_dotenv

# Load the API key from .env file
load_dotenv()
openai.api_key = os.environ["OPENAI_API_KEY"]


def generate_text(prompt, retries=3):
    for i in range(retries):
        try:
            response = openai.Completion.create(
                engine="text-davinci-002",
                prompt=prompt,
                max_tokens=200, # Try to keep this below 200
                n=1,
                stop=None,
                temperature=0.5,
            )

            return response.choices[0].text.strip()

        except openai.error.APIError as e:
            print(f"OpenAI API error ({e.status}): {e.message}")
            if i == retries - 1:
                raise e
            wait_time = 2 ** i
            print(f"Retrying in {wait_time} seconds...")
            time.sleep(wait_time)


def save_to_doc_file(filename, title, content, directory="output"):
    doc = docx.Document()

    # Add title
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    doc.add_paragraph()

    # Add bullet points
    for line in content.split('\n'):
        bullet_paragraph = doc.add_paragraph(style='ListBullet')
        bullet_run = bullet_paragraph.add_run(line)
        bullet_run.font.size = Pt(12)

    # Directory to save the file to
    filename = os.path.join(directory, filename)

    # Check if the directory exists, create it if it doesn't
    if not os.path.exists(directory):
        os.mkdir(directory)

    # Save to docx file. Create numbered file if file already exists
    if os.path.exists(filename):
        i = 1
        while os.path.exists(f"{filename[:-5]}_{i}.docx"):
            i += 1
        doc.save(f"{filename[:-5]}_{i}.docx")
    else:
        doc.save(filename)


if __name__ == "__main__":
    # Get user input
    topic = input("Enter a topic (default: 'Cars'): ") or 'Cars'
    adjective = input(
        "Enter an adjective (default: 'interesting'): ") or 'interesting'
    num_facts = input("Enter the number of facts (default: 10): ") or '10'
    format_ = input(
        "Enter the format (default: 'list'): ") or 'list'
    importance = input(
        "Enter the importance (default: 'random'): ") or 'random'
    category = input("Enter the category (default: 'None'): ")
    simplify = input("Simplify the text? (y/n) (default: n): ") or 'n'

    # Generate text
    prompt = f"Write {num_facts} {adjective} {importance} facts about {topic} "
    if category:
        prompt += f"in the category: {category} "
    if format_:
        prompt += f"in the format: {format_} "
    if simplify == "y":
        prompt += "\nSimplify the text."
    else:
        prompt += "\nDon't simplify the text."

    generated_text = generate_text(prompt)

    # Remove blank lines
    lines = [line for line in generated_text.splitlines() if line.strip()
             ]  # filter out blank lines
    result = "\n".join(lines)

    # Save to doc file
    # Create a separate docx file for the topic
    output_filename = f"{topic.replace(' ', '_').lower()}_facts.docx"

    # Save generated text to docx file
    save_to_doc_file(
        output_filename, f"{num_facts} {adjective.title()} {importance.title()} Facts About {topic.title()} ({category.title()})", result)

    # Print results
    print(colored(f"\nGenerated text saved to {output_filename}", "green"))
    print("Prompt:\n")
    print(colored(prompt, "yellow"))
    print("\n")
    print("Generated content:\n")
    print(colored(generated_text, "cyan"))
